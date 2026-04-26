#!/usr/bin/env python3
"""Generate small PDF + DOCX fixtures for the document-attachment E2E test.

Run once when the corpus needs to change:
    pip install --user reportlab python-docx
    python3 testdata/documents/generate_fixtures.py

The output is checked into the repo so the E2E test does not require
reportlab/python-docx at runtime — only when the fixtures need updating.

PII corpus is intentionally German-leaning so the Presidio scanner
(German recognizers in `ingestion/pii_config.yaml`) flags PERSON +
EMAIL_ADDRESS reliably.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

OUT_DIR = Path(__file__).resolve().parent

PII_TEXT = (
    "Herr Sebastian Müller hat das Projekt Alpha am 15.03.2026 "
    "abgeschlossen. Kontakt: sebastian.mueller@example.com."
)


def _write_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    c.setFont("Helvetica", 12)
    c.drawString(72, 800, "Powerbrain E2E document fixture")
    c.drawString(72, 780, PII_TEXT)
    c.showPage()
    c.save()


def _write_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Powerbrain E2E document fixture", level=1)
    doc.add_paragraph(PII_TEXT)
    doc.save(str(path))


def main() -> None:
    pdf_path = OUT_DIR / "sample_with_pii.pdf"
    docx_path = OUT_DIR / "sample_with_pii.docx"
    _write_pdf(pdf_path)
    _write_docx(docx_path)
    print(f"Wrote {pdf_path} ({pdf_path.stat().st_size} bytes)")
    print(f"Wrote {docx_path} ({docx_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
